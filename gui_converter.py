
import tkinter as tk
from tkinter import filedialog, messagebox, ttk, scrolledtext
from docx import Document
from PIL import Image
import fitz
import os
from moviepy import VideoFileClip
import subprocess

# var for out put dir
output_directory = os.getcwd()

# groups and stuff
audio_formats = ["mp3", "wav", "aac", "flac", "wma", "m4a"]
video_formats = ["mp4", "mov", "mpg", "webm"]
image_formats = ["png", "jpg", "jpeg", "webp", "bmp"]
document_formats = ["pdf", "docx"]

supported_pairs = set()

# valid combos
for v in video_formats:
    for a in audio_formats:
        supported_pairs.add((v, a))
for img_in in image_formats:
    for img_out in image_formats:
        if img_in != img_out:
            supported_pairs.add((img_in, img_out))
supported_pairs.add(("pdf", "docx"))
supported_pairs.add(("docx", "pdf"))

def open_file(filetypes):
    file_path = filedialog.askopenfilename(filetypes=filetypes)
    if not file_path:
        messagebox.showwarning("No File", "No file selected.")
    return file_path

def choose_output_directory():
    global output_directory
    output_directory = filedialog.askdirectory()
    if output_directory:
        update_status(f"Output folder set to: {output_directory}")

def update_status(message):
    status_bar.config(text=message)
    log_output.insert(tk.END, message + "\n")
    log_output.see(tk.END)
    window.update_idletasks()
# CHANGE THIS NOW     744
def convert_file():
    try:
        input_format = input_format_var.get().lower()
        output_format = output_format_var.get().lower()

        if input_format == output_format:
            messagebox.showinfo("Invalid Selection", "Please select different input and output formats.")
            return

        if (input_format, output_format) not in supported_pairs:
            messagebox.showerror("Unsupported Conversion", f"Conversion from {input_format.upper()} to {output_format.upper()} is not supported.")
            return

        filetypes = [(f"{input_format.upper()} files", f"*.{input_format.lower()}")]
        input_path = open_file(filetypes)
        if not input_path:
            return

        base_name = os.path.splitext(os.path.basename(input_path))[0]
        output_path = os.path.join(output_directory, f"{base_name}.{output_format}")

        update_status(f"Converting {input_format.upper()} to {output_format.upper()}...")

        if input_format == "pdf" and output_format == "docx":
            doc = Document()
            os.makedirs("pdf_images", exist_ok=True)
            pdf = fitz.open(input_path)
            for page_num, page in enumerate(pdf, start=1):
                text = page.get_text("text")
                if text:
                    for line in text.split('\n'):
                        doc.add_paragraph(line)

                for img_index, img in enumerate(page.get_images(full=True), start=1):
                    xref = img[0]
                    base_image = pdf.extract_image(xref)
                    image_bytes = base_image["image"]
                    image_ext = base_image["ext"]

                    image_filename = os.path.join("pdf_images", f"page{page_num}_img{img_index}.{image_ext}")
                    with open(image_filename, "wb") as img_file:
                        img_file.write(image_bytes)

                    doc.add_picture(image_filename)
            doc.save(output_path)
            update_status(f"PDF converted to DOCX and saved as {output_path}")
# MORE FORMAT
        elif input_format == "docx" and output_format == "pdf":
            doc = Document(input_path)
            pdf_writer = fitz.open()
            for para in doc.paragraphs:
                pdf_writer.new_page().insert_text((72, 72), para.text)
            pdf_writer.save(output_path)
            update_status(f"DOCX converted to PDF and saved as {output_path}")

        elif input_format in video_formats and output_format in audio_formats:
            clip = VideoFileClip(input_path)
            clip.audio.write_audiofile(output_path)

        elif input_format in image_formats and output_format in image_formats:
            image = Image.open(input_path)
            image.save(output_path)

        else:
            subprocess.run(["ffmpeg", "-i", input_path, output_path], check=True)

        update_status(f"Conversion complete! Saved as {output_path}")

    except subprocess.CalledProcessError as e:
        update_status("Error occurred during conversion.")
        messagebox.showerror("Error", f"ffmpeg failed: {e}")
    except Exception as e:
        update_status("Error occurred during conversion.")
        messagebox.showerror("Error", f"An error occurred: {str(e)}")

def toggle_dark_mode():
    dark = window.tk.call("ttk::style", "theme", "use") == "default"
    if dark:
        window.tk.call("ttk::style", "theme", "use", "clam")
        window.configure(bg="#2e2e2e")
        for widget in window.winfo_children():
            try:
                widget.configure(bg="#2e2e2e", fg="white")
            except:
                pass
        update_status("Dark mode enabled")
    else:
        window.tk.call("ttk::style", "theme", "use", "default")
        window.configure(bg="SystemButtonFace")
        for widget in window.winfo_children():
            try:
                widget.configure(bg="SystemButtonFace", fg="black")
            except:
                pass
        update_status("Dark mode disabled")

# GUI STUFF
window = tk.Tk()
window.title("File Converter")
window.geometry("600x500")

frame = tk.Frame(window)
frame.pack(pady=20)

input_label = tk.Label(frame, text="Input Format:")
input_label.grid(row=0, column=0, padx=5, pady=5, sticky="e")

input_format_var = tk.StringVar()
input_dropdown = ttk.Combobox(frame, textvariable=input_format_var)
input_dropdown['values'] = audio_formats + video_formats + image_formats + document_formats
input_dropdown.grid(row=0, column=1, padx=5, pady=5)

output_label = tk.Label(frame, text="Output Format:")
output_label.grid(row=1, column=0, padx=5, pady=5, sticky="e")

output_format_var = tk.StringVar()
output_dropdown = ttk.Combobox(frame, textvariable=output_format_var)
output_dropdown['values'] = audio_formats + video_formats + image_formats + document_formats
output_dropdown.grid(row=1, column=1, padx=5, pady=5)

convert_button = tk.Button(window, text="Convert File", command=convert_file, height=2, width=20)
convert_button.pack(pady=10)

destination_button = tk.Button(window, text="Choose Output Folder", command=choose_output_directory)
destination_button.pack(pady=5)

dark_mode_button = tk.Button(window, text="Toggle Dark Mode", command=toggle_dark_mode)
dark_mode_button.pack(pady=5)

log_output = scrolledtext.ScrolledText(window, height=6)
log_output.pack(padx=10, pady=10, fill=tk.BOTH, expand=True)

status_bar = tk.Label(window, text="Ready", bd=1, relief=tk.SUNKEN, anchor="w")
status_bar.pack(side=tk.BOTTOM, fill=tk.X)

window.mainloop()
